import os
import datetime
import logging
from complementares.extrair_bioma import extrair_palavra, extrair_item_anterior
from extrair_gpt import criar_documento
from complementares.manipular_documento import substituir_textos_docx, gerar_nome_arquivo_unico
from calculo import realizar_calculo
from complementares.extenso import numero_completo_por_extenso
from complementares.extrair_data import extracao_anos
from complementares.dados_manager import DadosManager
from docx import Document
from docx.shared import Pt
from extrair_nome import extrair_linha_por_referencia
# from complementares.palavra_anterior import extrair_item_anterior

# Inicializando o gerenciador de dados
dados_manager = DadosManager()

def ajustar_formato_documento(doc):
    """Ajusta o formato de fontes no documento para Arial, 12pt."""
    try:
        for p in doc.paragraphs:
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
    except Exception as e:
        logging.error(f'Erro ao ajustar o formato do documento: {str(e)}')

def substituir_textos_no_documento(caminho_tac, dicionario, output_path):
    """Substitui os textos em um arquivo DOCX usando um dicionário de termos e ajusta o formato."""
    try:
        doc = Document(caminho_tac)
        for key, value in dicionario.items():
            for p in doc.paragraphs:
                if key in p.text:
                    p.text = p.text.replace(key, str(value))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
        ajustar_formato_documento(doc)
        doc.save(output_path)
    except Exception as e:
        logging.error(f'Erro ao substituir textos no documento: {str(e)}')

def calcular_anos_meses(data_extraida):
    """Calcula a diferença entre a data extraída e a data atual, retornando anos e meses."""
    data_atual = datetime.date.today()
    
    anos_diferenca = data_atual.year - data_extraida.year
    meses_diferenca = data_atual.month - data_extraida.month
    dias_diferenca = data_atual.day - data_extraida.day

    # Ajusta a diferença se necessário
    if meses_diferenca < 0 or (meses_diferenca == 0 and dias_diferenca < 0):
        anos_diferenca -= 1
        meses_diferenca += 12
        
    return float(str(anos_diferenca) + '.' + str(meses_diferenca))

def processar_arquivo_alerta(caminho_arquivo):
    """Processa o arquivo de alerta, extraindo o bioma e calculando a diferença de tempo."""
    try:
        bioma = extrair_palavra(caminho_arquivo, 'BIOMAS', 2).replace('MUNICÍPIO', '').strip()
        dados_manager.set_bioma(bioma)

        data_extraida = extracao_anos(caminho_arquivo)
        dados_manager.set_data(data_extraida)

        anos_meses = calcular_anos_meses(data_extraida)
        dados_manager.set_anos(anos_meses)

        logging.info(f'Bioma extraído: {dados_manager.get_bioma()}')
        logging.info(f'Diferença em anos e meses: {anos_meses}')
    except Exception as e:
        logging.error(f'Erro ao processar arquivo de alerta: {str(e)}')

def criar_dicionario_tac(caminho_arquivo):
    """Cria o dicionário TAC com informações extraídas e calculadas do arquivo."""
    try:
        dicionario = criar_documento(caminho_arquivo)
        bioma = dados_manager.get_bioma()
        area_afetada = dicionario.get('$area')
        n1 = dados_manager.get_anos()
        ano = dados_manager.get_data().year

        dicionario['$CEFIR'] = extrair_item_anterior('2.14. “Relatório CEFIR” (Anexo', caminho_arquivo)
        valor = dicionario['$valor'] = realizar_calculo(bioma, area_afetada, n1, ano)
        dicionario['$extenso'] = numero_completo_por_extenso(valor)
        dicionario['$proprietario'] = extrair_linha_por_referencia(caminho_arquivo, '2.4. Requerente do cadastro')

        return dicionario
    except Exception as e:
        logging.error(f'Erro ao criar dicionário TAC: {str(e)}')
        return None

def processar_arquivo_relatorio(caminho_arquivo, caminho_tac_lista):
    """
    Processa o relatório TAC, substituindo os valores no documento e salvando a saída.
    
    Parâmetros:
        caminho_arquivo (str): Caminho do arquivo de entrada.
        caminho_tac_lista (list): Lista de caminhos para os arquivos TAC.
    """
    try:
        # Tenta criar o dicionário TAC
        dicionario = None
        while dicionario is None:
            try:
                dicionario = criar_dicionario_tac(caminho_arquivo)
            except Exception as e:
                logging.error(f'Erro ao criar dicionário TAC: {str(e)}')
                print(f"Tentando novamente criar o dicionário TAC devido ao erro: {str(e)}")

        # Processa cada caminho TAC na lista
        for caminho_tac in caminho_tac_lista:
            sucesso = False
            print(caminho_arquivo)
            while not sucesso:
                try:
                    if dicionario:
                        output_path = gerar_caminho_saida(caminho_arquivo, caminho_tac)
                        substituir_textos_no_documento(caminho_tac, dicionario, output_path)
                        logging.info(f'Relatório processado e salvo em: {output_path}')
                        sucesso = True
                except ValueError as ve:
                    logging.error(f'Valor inválido encontrado ao processar {caminho_tac}: {str(ve)}')
                    print(f"Erro de valor inválido ao processar {caminho_tac}: {str(ve)}")
                    sucesso = True  # Continua, já que o erro foi tratado
                except Exception as e:
                    logging.error(f'Erro ao processar arquivo de relatório {caminho_tac}: {str(e)}')
                    print(f"Tentando novamente o arquivo {caminho_tac} devido ao erro: {str(e)}")
                    # O loop continuará até que o sucesso seja True
    except Exception as e:
        logging.error(f'Erro inesperado ao processar relatório: {str(e)}')
        print(f"Erro inesperado ao processar relatório: {str(e)}")


def gerar_caminho_saida(caminho_arquivo, caminho_tac):
    """Gera o nome do arquivo de saída com base no modelo de TAC, sem depender de 'Alerta'."""
    try:
        # Extrai a pasta de saída
        pasta_saida = os.path.dirname(caminho_arquivo)

        # Tenta extrair um número de alerta baseado no nome da pasta ou do arquivo
        nome_pasta = os.path.basename(pasta_saida)
        nome_arquivo = os.path.basename(caminho_arquivo)

        # Procura um número no nome do arquivo ou da pasta
        numeros_alerta = ''.join(filter(str.isdigit, nome_arquivo)) or ''.join(filter(str.isdigit, nome_pasta))

        # Se não houver número, usa a data atual como um fallback
        if not numeros_alerta:
            numeros_alerta = datetime.datetime.now().strftime("%Y%m%d")

        # Define o modelo baseado no conteúdo do caminho TAC
        modelo = 'COMPENSAÇÃO' if 'COMPENSAÇÃO' in caminho_tac else 'RECUPERAÇÃO'
        output_base_name = f'TAC_{numeros_alerta} - {modelo}'
        output_ext = '.docx'

        # Gera o caminho de saída usando a função auxiliar gerar_nome_arquivo_unico
        output_path = gerar_nome_arquivo_unico(output_base_name, output_ext, pasta_saida)
        
        return output_path
    except Exception as e:
        logging.error(f'Erro ao gerar caminho de saída: {str(e)}')
        raise

def processar_arquivo(caminho_arquivo, caminho_tac_lista):
    """Determina o tipo de arquivo e o processa adequadamente."""
    try:
        nome_arquivo = os.path.basename(caminho_arquivo)
        if nome_arquivo.startswith('ANEXO 01'):
            processar_arquivo_alerta(caminho_arquivo)
        elif nome_arquivo.startswith('Relatório'):
            processar_arquivo_relatorio(caminho_arquivo, caminho_tac_lista)
    except Exception as e:
        logging.error(f'Erro ao processar arquivo: {str(e)}')
